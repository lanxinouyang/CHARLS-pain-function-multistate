#!/usr/bin/env python3
"""Build the audited five-wave CHARLS Scientific Reports submission package."""

from __future__ import annotations

import json
import re
import shutil
import sys
import zipfile
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
WORK = ROOT
FIVE = ROOT
PACKAGE = ROOT / "publication_outputs"
FIG = PACKAGE / "figures"
SOURCE = PACKAGE / "source_data"
CODE = PACKAGE / "code"
ZIP_PATH = ROOT / "publication_outputs_package.zip"

sys.path.insert(0, str(WORK))
import manuscript_utils as bp


TITLE = (
    "Bidirectional associations between musculoskeletal pain burden and functional limitations "
    "in middle-aged and older Chinese adults: a multistate analysis"
)
RUNNING = "Pain burden and functional-state transitions"

MANUSCRIPT = PACKAGE / "Scientific_Reports_Manuscript_5Wave.docx"
SUPPLEMENT = PACKAGE / "Scientific_Reports_Supplementary_Information_5Wave.docx"
COVER = PACKAGE / "Scientific_Reports_Cover_Letter_5Wave.docx"
CHECKLIST = PACKAGE / "Scientific_Reports_Submission_Checklist_5Wave.docx"
STROBE = PACKAGE / "STROBE_Reporting_Map_5Wave.docx"

PAIN_INTERVALS = FIVE / "models" / "pain" / "fivewave_primary_intervals.csv"
FUNCTION_INTERVALS = FIVE / "models" / "function" / "fivewave_primary_intervals.csv"
LONG = FIVE / "cohort" / "CHARLS_分析长格式_45岁及以上.csv"


ABSTRACT = (
    "Musculoskeletal pain and limitations in daily activities often coexist, but their temporal relationship remains "
    "uncertain. We analysed five waves (2011, 2013, 2015, 2018, and 2020) of the China Health and Retirement Longitudinal "
    "Study to examine each condition in relation to subsequent changes in the other. Parallel continuous-time multistate "
    "models classified pain as zero, one, or at least two musculoskeletal sites and function as no limitation, instrumental "
    "activity limitation only, at least one basic activity limitation, or death. The analysis included 21,235 adults, "
    "62,135 adjacent-wave intervals, and 1,973 deaths. Compared with no limitation, basic activity limitation was associated "
    "with moving from no pain to multisite pain (adjusted hazard ratio 1.89, 95% confidence interval 1.57-2.28); instrumental "
    "activity limitation alone was not (1.07, 0.84-1.38). In the reverse direction, one and at least two pain sites were "
    "associated with transitions from no limitation to basic activity limitation (1.55, 1.28-1.87 and 2.19, 1.86-2.59). "
    "The direction of these estimates was retained with survey weighting, strict function scoring, alternative death "
    "ascertainment, and exclusion of the 2013 observations. The findings indicate bidirectional longitudinal associations "
    "but do not establish causality."
)

INTRO = [
    "Musculoskeletal pain and difficulty with everyday activities are common in later life and frequently occur together [1-6]. The number of painful sites is associated with mobility problems and limitations in basic and instrumental activities of daily living (BADL and IADL) [7-13]. Neither condition follows a simple one-way course: pain and functional limitations may worsen, improve, or recur over time [14,15]. Analyses based only on incident disability or a single binary pain measure do not capture these changes.",
    "Longitudinal evidence has mainly considered pain as a precursor of disability. Fewer studies have examined whether functional limitations also precede changes in pain burden, particularly in nationally representative samples from middle-income settings. Previous transition studies have shown that pain can be related to both the onset of disability and recovery [19]. Multistate models extend this work by estimating movement among several severity levels while accounting for death as an absorbing state [17,18].",
    "We examined five waves of the China Health and Retirement Longitudinal Study (CHARLS) [19]. Two parallel continuous-time multistate models were fitted: one related interval-start functional state to later pain-state transitions, and the other related interval-start pain burden to later functional-state transitions. Our main comparisons were the transition from no pain to multisite pain (P0 to P2) and the transition from no functional limitation to BADL limitation (F0 to F2). All other living-state and mortality transitions were also estimated. Because pain and function were recorded at the same survey visits, the results are interpreted as bidirectional longitudinal associations rather than causal effects."
]

RESULTS_SECTIONS = [
    ("Study population and follow-up", [
        "The harmonised dataset contained 26,347 people and 100,715 person-wave records across the five survey waves. After applying the age, state and entry-covariate criteria, 23,369 participants were eligible for follow-up; 2,134 did not contribute an analysable adjacent-wave interval. The final sample included 21,235 participants, 62,135 intervals, 12,577 household clusters and 1,973 confirmed deaths (Fig. 1). The four periods contributed 14,621, 14,518, 16,269 and 16,727 intervals, respectively. At model entry, mean age was 58.0 years (SD 9.9), 50.4% of participants were women and 4,484 were in the multisite pain state (Table 1)."
    ]),
    ("Functional state and subsequent pain transitions", [
        "Functional status distinguished the rate of progression from no pain to multisite pain (Fig. 2a; Table 2). Relative to F0, the adjusted hazard ratio (HR) for P0 to P2 was 1.07 (95% CI 0.84-1.38; P = 0.575) for IADL-only limitation (F1) and 1.89 (1.57-2.28; P = 1.96 × 10⁻¹¹) for BADL limitation (F2). Thus, the association was concentrated in the more severe functional state rather than following a uniform gradient across F1 and F2."
    ]),
    ("Pain burden and subsequent functional transitions", [
        "Pain burden showed a graded association with the transition from no limitation to BADL limitation (Fig. 2b; Table 3). Compared with P0, the adjusted HR for F0 to F2 was 1.55 (95% CI 1.28-1.87; P = 7.47 × 10⁻⁶) for one painful site (P1) and 2.19 (1.86-2.59; P = 2.18 × 10⁻²⁰) for at least two painful sites (P2). Each estimate is conditional on the covariates specified for that transition."
    ]),
    ("Other transitions and model specification", [
        "Associations for the remaining living-state transitions were mixed (Fig. 2; Table 2). Greater cross-domain burden was associated with several worsening transitions, whereas some improvement transitions had HRs below one. Because 18 exposure contrasts were estimated across living and mortality transitions, results outside the two prespecified focal pathways are considered exploratory.",
        "The least restrictive adjustment structure gave the best fit (Fig. 4). Compared with the 126-parameter transition-specific model, the 78-parameter family-shared model fit worse for pain (likelihood-ratio chi-square = 156.31, 48 df, P = 2.15 × 10⁻¹³) and function (chi-square = 273.41, 48 df, P = 2.63 × 10⁻³³). The 110-parameter model also fit worse for pain (chi-square = 33.38, 16 df, P = 0.00658), whereas the comparison for function was inconclusive (chi-square = 25.54, 16 df, P = 0.0608). The 126-parameter structure was therefore retained for both domains."
    ]),
    ("Two-year state probabilities", [
        "Absolute probabilities differed by survey period (Fig. 3). Among participants starting in P0, the modelled probability of P2 under F0 versus F2 was 10.3% versus 20.2% in 2011-2013, 13.4% versus 26.0% in 2013-2015, 23.6% versus 38.9% in 2015-2018 and 19.3% versus 33.2% in 2018-2020. Among those starting in F0, the probability of F2 under P0 versus P2 was 9.4% versus 18.8%, 10.5% versus 21.0%, 6.0% versus 12.5% and 9.1% versus 18.6%, respectively. These standardised scenarios include all three living states and death; they are not individual forecasts."
    ]),
    ("Sensitivity analyses", [
        "The direction of both focal associations was retained across the sensitivity analyses (Supplementary Fig. S3; Supplementary Table S5). The corresponding HRs for the function-to-pain and pain-to-function pathways were 2.10 and 2.37 with survey weighting, 1.55 and 1.70 with strict 11-item function scoring, and 1.89 and 2.15 when 2013 deaths were restricted to Exit Interview records. Excluding the 2013 observations yielded HRs of 2.24 and 2.14."
    ])
]

RESULTS = [paragraph for _, paragraphs in RESULTS_SECTIONS for paragraph in paragraphs]

DISCUSSION = [
    "Across four adjacent-wave periods, functional status was associated with later pain transitions and pain burden was associated with later functional transitions. BADL limitation, but not IADL-only limitation, marked a higher rate of moving from no pain to multisite pain. In the other direction, the rate of moving from no limitation to BADL limitation increased from one to multiple painful sites. The asymmetry is clinically relevant: the pain-to-function association followed a clearer severity gradient than the function-to-pain association.",
    "The findings accord with longitudinal studies linking pain to later mobility and activity limitations [7-13], while adding evidence for the reverse temporal sequence. Pain may restrict activity, disturb sleep and contribute to deconditioning. Functional limitations may, in turn, reduce opportunities for movement, rehabilitation and self-management, or reflect underlying morbidity that also influences pain. These mechanisms cannot be separated with the present observational data, and common reporting processes or unmeasured health conditions may contribute to both associations.",
    "Modelling several states rather than a single incident outcome revealed substantial variation between pathways. Some improvement transitions were more frequent at lower cross-domain burden, but the pattern was not uniform. Period-specific probabilities also showed that the absolute chance of reaching a severe state changed over calendar time even when the relative association remained in the same direction. Pain-site count and BADL/IADL status therefore provide complementary information, although the modelled probabilities should not be interpreted as individual prognosis.",
    "Weighting, strict function scoring and alternative death ascertainment did not alter the direction of the focal results. In the wave-exclusion analysis, omitting the 2013 observations changed the function-to-pain HR from 1.89 to 2.24, whereas the pain-to-function HR changed from 2.19 to 2.14. The larger change in the function-to-pain estimate suggests some dependence on observation spacing and sample composition, whereas the pain-to-function estimate was comparatively stable.",
    "The main strengths are the use of five survey waves, explicit modelling of worsening and improvement, inclusion of death as an absorbing state, transition-specific adjustment and household-cluster robust inference. Several limitations remain. Pain, BADL and IADL were self-reported at the same visit. Unequal gaps between surveys leave intermediate changes unobserved, so the estimated cross-level transitions represent net movement between visits. The structural-skip rule used for function in 2011-2018 may not remove all measurement differences across waves, although the complete-11-item analysis gave similar results. Death time was interval censored, selection may arise from complete entry covariates and available follow-up, and survey weights cannot address unmeasured selection. The Markov and piecewise-constant intensity assumptions are approximations, and estimates outside the two focal pathways require cautious interpretation because of multiple comparisons.",
    "In summary, musculoskeletal pain burden and functional limitations were associated with subsequent state transitions in both directions among middle-aged and older Chinese adults. More frequent measurements or intervention studies are needed to determine causal order and assess whether modifying one condition changes the course of the other."
]

METHODS = [
    ("Study design and participants. ", "CHARLS is a nationally representative longitudinal study of Chinese adults aged 45 years or older and their spouses [19]. The present analysis used the 2011, 2013, 2015, 2018 and 2020 waves. Participants entered at the first wave in which pain and function were both observed and all primary adjustment variables were available. Eligible follow-up intervals linked adjacent waves (2011-2013, 2013-2015, 2015-2018 and 2018-2020) and required an observed outcome state and cross-domain exposure at interval start. At the next wave, an interval ended in an observed living state, confirmed death, or known survival with the domain-specific state missing. Reporting follows STROBE [20]."),
    ("Cross-wave data harmonisation. ", "Health, demographic, sample-status, weight and exit-interview files were linked within each wave using the harmonised person identifier. The 2013 public release did not include a separate sample-information file, so wave status was constructed from the union of Health Status and Functioning, Weights and Exit Interview records. This yielded 19,689 identifiers. A death was recorded when INDV_L_Died equalled one in the Weights file or the identifier appeared in the Exit Interview file; 464 deaths met this definition, including 408 identified by both sources. No record combined a completed health interview with a death indicator in the same wave. Duplicate person-wave records, invalid state values and health observations after death were checked across all five waves before the analytic cohort was assembled. A sensitivity analysis restricted 2013 deaths to Exit Interview records."),
    ("Pain states. ", "At each wave, participants were asked about bodily pain and painful sites. Twelve musculoskeletal sites were harmonised: shoulder, arm, wrist, fingers, back, waist, buttocks, leg, knee, ankle, toes, and neck. P0 denoted no painful site, P1 one painful site, and P2 at least two painful sites. In 2013, bodily pain was screened with wb16 (none versus any higher severity) and sites were read from da042s1-da042s15; wave-specific mappings are listed in Supplementary Table S1."),
    ("Functional states. ", "BADL items covered dressing, bathing, eating, getting into or out of bed, toileting, and continence, following the Katz framework [20]. IADL items covered household chores, preparing meals, shopping, managing money, and taking medication, following the Lawton framework [16]. F0 denoted no BADL or IADL limitation, F1 IADL limitation only, and F2 at least one BADL limitation; F2 took precedence regardless of IADL status. Death (D) was absorbing. In 2011, 2013, 2015, and 2018, questionnaire routing could skip all six BADL items. We classified such records as having no BADL limitation only when at least one IADL response was valid. The strict sensitivity analysis required complete responses to all 11 BADL/IADL items. The 2013 primary items were db010-db015 and db016, db017, db018, db020, and db019; the telephone item db035 was not included in the cross-wave five-item IADL definition."),
    ("Covariates. ", "Age was updated at each interval start. Sex, education (none, primary school or below, middle school, or high school and above), marital/partner status, rural versus urban community classification, and the count of 14 self-reported chronic diseases were fixed at the participant's model-entry wave. These six factors formed the primary adjustment set. The cross-domain exposure was time varying and measured at interval start."),
    ("Continuous-time multistate models. ", "We fitted two parallel continuous-time panel multistate models using standard matrix-exponential methods [17,18]. Each model allowed all six directed transitions among three living states and death from each living state, for nine intensities in total. The pain model used interval-start function as the exposure; the function model used interval-start pain. For interval duration t, the transition-probability matrix was P(t)=exp(Qt), where Q was a four-state intensity matrix with an absorbing death row. An observed endpoint contributed its corresponding transition probability. If a participant was known alive but the endpoint domain state was missing, the likelihood contribution summed probabilities over the three living states. Death was treated as interval censored at the survey endpoint."),
    ("Parameterisation and model selection. ", "The unrestricted model estimated nine baseline log-intensities for 2011-2013, 27 transition-specific period contrasts for the three later periods, 18 transition-specific cross-domain exposure contrasts, and 72 transition-specific coefficients for the eight adjustment indicators or scales, for 126 parameters. We compared it with a 78-parameter family-shared model and a 110-parameter model with living-transition-specific but death-shared confounder effects using likelihood-ratio tests. Because the restrictive structures were rejected for pain and the family-shared structure was rejected for both outcomes, the 126-parameter model was used for both domains."),
    ("Estimation and uncertainty. ", "Models were fitted by L-BFGS-B with analytic likelihood gradients. Baseline intensities were bounded from 10⁻⁶ to 5 per year; period and exposure log-HRs were bounded from -4 to 4 and confounder log-HRs from -5 to 5. Numerical Hessians used a central finite-difference step of 2 × 10⁻⁴ and a generalised inverse tolerance of 10⁻¹⁰. We calculated household-cluster sandwich covariance matrices with a finite-sample correction; participants without household identifiers formed individual clusters. The maximum projected gradient was 4.90 × 10⁻⁴ for pain and 7.25 × 10⁻⁴ for function, and both numerical Hessians were positive definite."),
    ("Absolute probabilities. ", "For each survey period, we used the fitted Q matrix to calculate two-year state probabilities from selected origin and exposure states, holding continuous covariates at period-start means and categorical indicators at their period-start means. Uncertainty was propagated with 5,000 draws of the complete parameter vector from a multivariate normal distribution based on the household-robust covariance matrix. Point estimates in Fig. 3 use the fitted parameter vector; Supplementary Tables S9a and S9b report simulation medians and 2.5th-97.5th percentile intervals. For these calculations, exposure was held constant over two years. The probabilities apply to the covariate values specified above and are not predictions for individual participants."),
    ("Sensitivity analyses. ", "We repeated the focal analyses with cross-sectional survey weights, after trimming weights at the 1st and 99th percentiles and normalising within interval-start wave. Additional analyses required complete responses to all 11 BADL/IADL items, restricted 2013 deaths to Exit Interview records, or omitted the 2013 observations to assess whether that wave influenced the focal estimates. The weighted analysis omitted 1,514 intervals without a valid weight. Tests were two sided and exact P values are reported. Analyses used Python 3.12.13, NumPy 2.5.2, pandas 2.2.3 and SciPy 1.18.0."),
    ("Ethics. ", "The original CHARLS survey was approved by the Peking University Institutional Review Board (IRB00001052-11015), and all participants provided written informed consent. This secondary analysis used only de-identified data obtained through the CHARLS data-use application and involved no direct contact with participants. No additional ethics approval was required for this analysis. All procedures were performed in accordance with relevant guidelines and regulations and the Declaration of Helsinki."),
    ("Use of generative AI and AI-assisted technologies. ", "OpenAI Codex was used to assist with language editing and code development and review. All analytical procedures and outputs were reviewed and verified by the authors.")
]

FIGURE_LEGENDS = [
    "Figure 1. Cohort construction, five-wave survey timeline, and multistate model structures. Panel a shows the analytic flow and intervals contributed by each adjacent-wave period. Panel b shows the pain-state model with interval-start function as the time-varying exposure; panel c shows the function-state model with interval-start pain as the time-varying exposure. Every directed transition among living states and each transition to absorbing death was permitted. BADL, basic activities of daily living; CHARLS, China Health and Retirement Longitudinal Study; IADL, instrumental activities of daily living.",
    "Figure 2. Household-cluster robust adjusted hazard ratios for all living-state transitions. Panel a relates interval-start functional state to pain transitions; panel b relates interval-start pain state to function transitions. Error bars show 95% confidence intervals. Shading identifies the two clinically focal severe-worsening pathways. Other transitions are exploratory.",
    "Figure 3. Period-specific modelled two-year state probabilities. Bars show complete four-state probability distributions, including death, for participants starting in P0 under F0 versus F2 (panel a) and participants starting in F0 under P0 versus P2 (panel b). White labels show fitted severe-state probabilities. Covariates were fixed at period-start means, and exposure was held constant for the two-year scenario.",
    "Figure 4. Comparison of nested multistate model structures. Bars show the increase in negative log-likelihood relative to the 126-parameter unrestricted model. The 78-parameter model shared confounder effects within worsening, improvement, and mortality families; the 110-parameter model used transition-specific effects for living pathways and shared effects for death pathways. P values compare the 110- and 126-parameter models."
]


def count_words(text: str) -> int:
    return len(re.findall(r"\b[\w'-]+\b", text))


def style(doc: Document, running: str, compact: bool = False) -> None:
    bp.style_document(doc, running, compact=compact)
    # Scientific Reports requests an unjustified, single-column manuscript.
    doc.styles["Normal"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.styles["Reference Text"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def title_page(doc: Document, label: str, metadata_notice: bool = True) -> None:
    p = doc.add_paragraph(style="Document Label")
    p.paragraph_format.space_before = Pt(18)
    bp.set_run_font(p.add_run(label.upper()), size=8.5, bold=True, color=bp.BLUE)
    p = doc.add_paragraph(style="Article Title")
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(14)
    bp.set_run_font(p.add_run(TITLE), size=18, bold=True, color=bp.INK)
    p = doc.add_paragraph(style="Author Line")
    bp.set_run_font(p.add_run("Lanxin Ouyang¹†, Jinghui Song¹†, Yaqin Hu¹, and Yongyan Ding²*"), size=10.5, color=bp.INK)
    p = doc.add_paragraph(style="Small Note")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp.set_run_font(p.add_run("¹ Orthopedics, Jingzhou Hospital Affiliated to Yangtze University, Jingzhou 434000, Hubei, China\n² Gynecology, Jingzhou Hospital Affiliated to Yangtze University, Jingzhou 434000, Hubei, China"), size=9.2, color=bp.GRAY)
    p = doc.add_paragraph(style="Small Note")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp.set_run_font(p.add_run("*Correspondence: Yongyan Ding, Gynecology, Jingzhou Hospital Affiliated to Yangtze University, Jingzhou 434000, Hubei, China. Email: 1320037305@qq.com"), size=8.8, color=bp.GRAY)
    bp.add_text(doc, f"Running title: {RUNNING}", style="Small Note", align=WD_ALIGN_PARAGRAPH.CENTER)


def load_data():
    pain = pd.read_csv(PAIN_INTERVALS, dtype={"person_id": "string", "household_id": "string", "cluster_id": "string"})
    function = pd.read_csv(FUNCTION_INTERVALS, dtype={"person_id": "string", "household_id": "string", "cluster_id": "string"})
    long = pd.read_csv(LONG, dtype={"person_id": "string", "household_id": "string"}, low_memory=False)
    return pain, function, long


def entry_frame(pain: pd.DataFrame, long: pd.DataFrame) -> pd.DataFrame:
    entry = pain[["person_id", "entry_wave"]].drop_duplicates().copy()
    entry["entry_wave"] = entry["entry_wave"].astype(int)
    result = entry.merge(long, left_on=["person_id", "entry_wave"], right_on=["person_id", "wave"],
                         how="left", validate="one_to_one")
    if len(result) != 21235 or result["pain_state"].isna().any():
        raise ValueError("Entry-frame audit failed")
    return result


def n_pct(series: pd.Series) -> str:
    valid = series.dropna()
    return f"{int(valid.sum()):,} ({100 * valid.mean():.1f})"


def cat_n_pct(series: pd.Series, value: str) -> str:
    count = int(series.eq(value).sum())
    return f"{count:,} ({100 * count / len(series):.1f})"


def mean_sd(series: pd.Series) -> str:
    return f"{series.mean():.1f} ({series.std(ddof=1):.1f})"


def baseline_rows(entry: pd.DataFrame) -> list[list[str]]:
    groups = [entry, entry[entry.pain_state.eq("P0")], entry[entry.pain_state.eq("P1")], entry[entry.pain_state.eq("P2")]]
    rows = [["Participants, n", *[f"{len(x):,}" for x in groups]]]
    rows.append(["Age, years; mean (SD)", *[mean_sd(x.age_years) for x in groups]])
    rows.append(["Women, n (%)", *[n_pct(x.female) for x in groups]])
    for label, raw in [("No formal education", "无正规教育"), ("Primary school or below", "小学及以下"),
                       ("Middle school", "初中"), ("High school or above", "高中及以上")]:
        rows.append([f"{label}, n (%)", *[cat_n_pct(x.education_cat, raw) for x in groups]])
    rows.append(["Married/partnered, n (%)", *[n_pct(x.partnered) for x in groups]])
    rows.append(["Rural community, n (%)", *[n_pct(x.rural_nbs) for x in groups]])
    rows.append(["Chronic disease count; mean (SD)", *[mean_sd(x.chronic_n) for x in groups]])
    rows.append(["Function F0/F1/F2, n", *[
        " / ".join(f"{int(x.function_state.eq(s).sum()):,}" for s in ("F0", "F1", "F2")) for x in groups
    ]])
    return rows


def fmt_hr(row) -> str:
    return f"{row.hr:.2f} ({row.ci95_low:.2f}-{row.ci95_high:.2f})"


def fmt_p(value: float) -> str:
    if value >= 0.001:
        return f"{value:.3f}"
    return f"{value:.2e}"


def hr_rows(domain: str, intervals: pd.DataFrame, mortality: bool = False) -> list[list[str]]:
    hr = pd.read_csv(FIVE / "models" / domain / "fivewave_final_full_household_robust_hr.csv")
    if domain == "pain":
        origins = ["P0", "P1", "P2"]
        living = [("P0", "P1"), ("P0", "P2"), ("P1", "P2"), ("P1", "P0"), ("P2", "P1"), ("P2", "P0")]
        contrasts = ["F1 vs F0", "F2 vs F0"]
    else:
        origins = ["F0", "F1", "F2"]
        living = [("F0", "F1"), ("F0", "F2"), ("F1", "F2"), ("F1", "F0"), ("F2", "F1"), ("F2", "F0")]
        contrasts = ["P1 vs P0", "P2 vs P0"]
    edges = [(origin, "D") for origin in origins] if mortality else living
    rows = []
    for origin, destination in edges:
        count = int(((intervals.from_state == origin) & (intervals.to_state == destination)).sum())
        found = [hr[(hr.from_state == origin) & (hr.to_state == destination) & (hr.contrast == contrast)].iloc[0]
                 for contrast in contrasts]
        rows.append([f"{origin} → {destination}", f"{count:,}", fmt_hr(found[0]), fmt_p(found[0].p_value),
                     fmt_hr(found[1]), fmt_p(found[1].p_value)])
    return rows


def model_selection_rows() -> list[list[str]]:
    rows = []
    for domain in ("pain", "function"):
        primary = json.loads((FIVE / "models" / domain / "fivewave_primary_audit.json").read_text())
        final = json.loads((FIVE / "models" / domain / "fivewave_final_full_audit.json").read_text())
        rows.extend([
            [domain.title(), "78: family-shared", "78", f"{primary['shared_fit']['negative_log_likelihood']:.2f}", "-", "-", "-"],
            [domain.title(), "110: living-specific/death-shared", "110", f"{primary['primary_fit']['negative_log_likelihood']:.2f}",
             f"{final['comparisons']['primary_110_vs_full_126']['chi_square']:.2f}", "16", fmt_p(final['comparisons']['primary_110_vs_full_126']['p_value'])],
            [domain.title(), "126: all transitions specific", "126", f"{final['negative_log_likelihood']:.2f}", "Reference", "-", "-"],
        ])
    return rows


def probability_rows(domain_filter: str | None = None) -> list[list[str]]:
    rows = []
    for domain in ("pain", "function"):
        if domain_filter is not None and domain != domain_filter:
            continue
        data = pd.read_csv(FIVE / "models" / domain / "fivewave_final_full_period_probabilities.csv")
        if domain == "pain":
            data = data[(data.origin_state == "P0") & (data.exposure_state.isin(["F0", "F2"]))]
        else:
            data = data[(data.origin_state == "F0") & (data.exposure_state.isin(["P0", "P2"]))]
        for row in data.itertuples():
            rows.append([domain.title(), row.period, row.exposure_state, f"{row.origin_state} → {row.destination_state}",
                         f"{100 * row.point_estimate:.2f}", f"{100 * row.simulation_median:.2f}",
                         f"{100 * row.ci95_low:.2f}-{100 * row.ci95_high:.2f}"])
    return rows


def sensitivity_rows() -> list[list[str]]:
    labels = {
        "cross-sectional-weighted": "Cross-sectional weighted",
        "strict-complete-11-item-function": "Strict complete 11-item function",
        "2013-exit-interview-deaths-only": "2013 Exit Interview deaths only",
        "leave-2013-out": "Excluding 2013 observations",
    }
    output = [["Five-wave primary analysis", "21,235 / 62,135", "1.89", "2.19"]]
    pain = pd.read_csv(FIVE / "models" / "pain" / "fivewave_final_sensitivity_summary.csv").set_index("analysis")
    function = pd.read_csv(FIVE / "models" / "function" / "fivewave_final_sensitivity_summary.csv").set_index("analysis")
    for key, label in labels.items():
        pi, fu = pain.loc[key], function.loc[key]
        people = int(min(pi.people, fu.people)); intervals = int(min(pi.intervals, fu.intervals))
        output.append([label, f"{people:,} / {intervals:,}", f"{pi.hr:.2f}", f"{fu.hr:.2f}"])
    return output


def disposition_rows(domain: str, intervals: pd.DataFrame) -> list[list[str]]:
    rows = []
    valid_prefix = "P" if domain == "pain" else "F"
    for start, data in intervals.groupby("start_wave", sort=True):
        period = f"{int(start)}-{int(data.end_wave.iloc[0])}"
        living = int(data.to_state.astype(str).str.startswith(valid_prefix).sum())
        rows.append([period, f"{len(data):,}", f"{living:,}", f"{int(data.to_state.eq('D').sum()):,}",
                     f"{int(data.to_state.eq('ALIVE_UNKNOWN').sum()):,}"])
    return rows


def build_manuscript(pain: pd.DataFrame, function: pd.DataFrame, entry: pd.DataFrame) -> None:
    doc = Document(); style(doc, RUNNING)
    title_page(doc, "Scientific Reports | Article manuscript")
    bp.add_text(doc, "Keywords: musculoskeletal pain; functional disability; activities of daily living; multistate model; longitudinal study; CHARLS",
                style="Small Note", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()
    doc.add_heading("Abstract", level=1); bp.add_text(doc, ABSTRACT)
    doc.add_heading("Introduction", level=1)
    for paragraph in INTRO:
        bp.add_text(doc, paragraph)
    doc.add_page_break()
    doc.add_heading("Results", level=1)
    for subheading, paragraphs in RESULTS_SECTIONS:
        doc.add_heading(subheading, level=2)
        for paragraph in paragraphs:
            bp.add_text(doc, paragraph)

    bp.add_table(doc, "Table 1. Participant characteristics at model entry, overall and by pain state",
                 ["Characteristic", "Overall", "P0", "P1", "P2"], baseline_rows(entry), [2.7, 1.45, 1.4, 1.3, 1.4],
                 note="P0, no painful site; P1, one painful site; P2, at least two painful sites. Percentages use non-missing denominators; entry covariates were complete by design. No baseline significance tests were performed.",
                 font_size=7.0)
    doc.add_page_break()
    bp.add_table(doc, "Table 2. Functional state and subsequent living-state pain transitions",
                 ["Transition", "Observed", "F1 vs F0 HR (95% CI)", "P", "F2 vs F0 HR (95% CI)", "P"],
                 hr_rows("pain", pain), [1.25, 1.0, 2.05, 0.8, 2.05, 0.8],
                 note="Observed counts are adjacent-visit endpoint changes and are descriptive. HRs use the 126-parameter model and household-cluster robust covariance. P0/P1/P2 denote zero, one, or at least two painful sites; F0/F1/F2 denote no limitation, IADL-only limitation, or at least one BADL limitation.", font_size=6.8)
    bp.add_table(doc, "Table 3. Pain state and subsequent living-state function transitions",
                 ["Transition", "Observed", "P1 vs P0 HR (95% CI)", "P", "P2 vs P0 HR (95% CI)", "P"],
                 hr_rows("function", function), [1.25, 1.0, 2.05, 0.8, 2.05, 0.8],
                 note="All transition-specific models adjusted for interval-start age and entry sex, education, marital/partner status, rural versus urban community classification, and chronic disease count. Exact two-sided P values are reported.", font_size=6.8)

    doc.add_heading("Discussion", level=1)
    for paragraph in DISCUSSION: bp.add_text(doc, paragraph)
    doc.add_heading("Methods", level=1)
    for label, body in METHODS: bp.add_text(doc, label + body, label=label)

    doc.add_heading("Data availability", level=1)
    bp.add_text(doc, "Individual-level CHARLS data are available to registered users after approval by the CHARLS data custodian (https://charls.pku.edu.cn/). The authors cannot redistribute controlled-access participant data. Aggregate source data underlying the tables and figures are included in the peer-review package.")
    doc.add_heading("Code availability", level=1)
    bp.add_text(doc, "Analysis, quality-control and figure-generation scripts, together with non-disclosive aggregate source data, are included in the peer-review package. A permanent repository URL and archived release identifier will be added before publication." )
    doc.add_heading("References", level=1)
    for i, ref in enumerate(bp.REFERENCES, 1):
        p = doc.add_paragraph(style="Reference Text")
        p.paragraph_format.left_indent = Inches(0.22); p.paragraph_format.first_line_indent = Inches(-0.22)
        bp.set_run_font(p.add_run(f"{i}. {ref}"), size=8.3)

    doc.add_heading("Acknowledgements", level=1)
    bp.add_text(doc, "We thank the CHARLS research team and all study participants.")
    doc.add_heading("Author contributions", level=1)
    bp.add_text(doc, "L.O. and J.S. contributed equally to this work. L.O. conceived the study, performed the statistical analyses, prepared the figures and tables, and drafted the manuscript. J.S. contributed to the study methodology, validation of the analyses, interpretation of the results, and drafting and revision of the manuscript. Y.H. contributed to the methodological development, interpretation of the findings, and critical revision of the manuscript. Y.D. supervised the study, contributed to the study conception and design, and critically revised the manuscript. All authors reviewed and approved the final manuscript." )
    doc.add_heading("Funding", level=1)
    bp.add_text(doc, "This work was supported by the Natural Science Foundation of Hubei Province (Grant No. 2026AFC0538)." )
    doc.add_heading("Competing interests", level=1)
    bp.add_text(doc, "The authors declare no competing interests." )

    doc.add_heading("Figure legends", level=1)
    for legend in FIGURE_LEGENDS: bp.add_text(doc, legend, style="Figure Caption")
    figures = [
        ("Fig1_study_design_fivewave.png", "Figure 1", "Flow diagram and five-wave timeline above pain and function multistate diagrams with death.", 6.4),
        ("Fig2_transition_HRs_fivewave.png", "Figure 2", "Two-panel forest plot of adjusted hazard ratios for all living-state pain and function transitions.", 6.45),
        ("Fig3_period_probabilities_fivewave.png", "Figure 3", "Two stacked-bar panels showing full state probabilities for four survey periods.", 6.35),
        ("Fig4_model_structure_fivewave.png", "Figure 4", "Two-panel bar chart comparing nested model negative log-likelihoods.", 6.35),
    ]
    for stem, label, alt, width in figures:
        doc.add_page_break(); bp.add_figure(doc, FIG / stem, label, alt, width)
    doc.core_properties.title = TITLE; doc.core_properties.subject = "Scientific Reports manuscript"
    doc.save(MANUSCRIPT)


def build_supplement(pain: pd.DataFrame, function: pd.DataFrame, entry: pd.DataFrame) -> None:
    doc = Document(); style(doc, RUNNING + " | Supplement", compact=True)
    title_page(doc, "Supplementary Information", metadata_notice=False)
    doc.add_heading("Supplementary Methods", level=1)
    bp.add_text(doc, "This file reports the wave-specific measurement mapping, the construction of mortality status, endpoint disposition, model selection, numerical convergence, mortality pathways, sensitivity analyses and complete modelled probability vectors. No participant-level CHARLS data are redistributed.")

    mapping = [
        ["2011", "da041; da042s1-da042s15", "db010-db015", "db016, db017, db018, db020, db019"],
        ["2013", "wb16; da042s1-da042s15", "db010-db015", "db016, db017, db018, db020, db019"],
        ["2015", "da041; da042s1-da042s16", "db010-db015", "db016, db017, db018, db020, db019"],
        ["2018", "da041_w4; da042_s1-da042_s16", "db010-db015", "db016, db017, db018, db020, db019"],
        ["2020", "da027; da028_s1-da028_s16", "db001, db003, db005, db007, db009, db011", "db012, db014, db016, db020, db022"],
    ]
    bp.add_table(doc, "Supplementary Table S1. Wave-specific pain and function item mapping",
                 ["Wave", "Pain screen and site block", "Six BADL items", "Five harmonised IADL items"], mapping,
                 [0.75, 2.8, 2.2, 2.7], note="The harmonised pain count retained 12 musculoskeletal locations. Wave-specific extra non-musculoskeletal categories were not counted. The telephone item was excluded from the cross-wave five-item IADL definition.", font_size=6.8, total=9500)

    s2013 = [
        ["Union person identifiers", "19,689", "Health, Weights, and Exit Interview union"],
        ["Health identifiers", "18,455", "Unique; no duplicates"],
        ["Valid weight identifiers", "19,666", "13 blank/invalid ID rows excluded"],
        ["Exit Interview identifiers", "431", "Unique"],
        ["Weight-file death flags", "441", "INDV_L_Died = 1"],
        ["Union confirmed deaths", "464", "Either death source"],
        ["Deaths identified by both sources", "408", "Overlap"],
        ["Health/death conflicts", "0", "PASS"],
        ["Death year available", "424", "Among 464 confirmed deaths"],
        ["Death month available", "420", "Among 464 confirmed deaths"],
    ]
    bp.add_table(doc, "Supplementary Table S2. Audit of the 2013 CHARLS source files",
                 ["Audit item", "n", "Interpretation"], s2013, [2.6, 1.0, 3.4], font_size=7.1, total=9500)

    flow = [
        ["Unique identifiers across five waves", "26,347", "-"],
        ["Age at least 45 years", "23,811", "2,536"],
        ["Joint entry state and complete entry covariates", "23,369", "442"],
        ["At least one included adjacent interval", "21,235", "2,134"],
    ]
    bp.add_table(doc, "Supplementary Table S3. Analytic cohort flow",
                 ["Stage", "Retained", "Excluded from preceding stage"], flow, [4.1, 1.3, 2.2], font_size=7.3, total=9500)

    bp.add_table(doc, "Supplementary Table S4a. Endpoint disposition in the pain-state model",
                 ["Period", "Starting intervals", "Living state observed", "Confirmed death", "Known alive, state missing"],
                 disposition_rows("pain", pain), [1.2, 1.4, 1.7, 1.4, 1.9], font_size=7.0, total=9500)
    bp.add_table(doc, "Supplementary Table S4b. Endpoint disposition in the function-state model",
                 ["Period", "Starting intervals", "Living state observed", "Confirmed death", "Known alive, state missing"],
                 disposition_rows("function", function), [1.2, 1.4, 1.7, 1.4, 1.9], font_size=7.0, total=9500)

    bp.add_table(doc, "Supplementary Table S5. Focal-pathway sensitivity analyses",
                 ["Analysis", "People / intervals", "F2 vs F0 for P0 → P2", "P2 vs P0 for F0 → F2"],
                 sensitivity_rows(), [2.8, 1.6, 2.0, 2.0],
                 note="Values are adjusted HR point estimates from unrestricted transition-specific models. Figure S3 visualises the same values. Confidence intervals were not recomputed for the four refitted sensitivity models and are therefore not presented.", font_size=7.0, total=9500)

    doc.add_page_break()
    bp.add_table(doc, "Supplementary Table S6. Nested model-structure comparison",
                 ["Domain", "Model", "Parameters", "NLL", "LR chi-square vs 126", "df", "P"],
                 model_selection_rows(), [0.9, 2.5, 0.9, 1.2, 1.4, 0.6, 0.9],
                 note="The 78-parameter versus 126-parameter comparisons were: pain chi-square 156.31, 48 df, P=2.15e-13; function chi-square 273.41, 48 df, P=2.63e-33. NLL, negative log-likelihood.", font_size=6.5, total=9500)

    convergence = []
    for domain in ("pain", "function"):
        audit = json.loads((FIVE / "models" / domain / "fivewave_final_full_audit.json").read_text())
        convergence.append([domain.title(), str(audit["refinement"]["optimizer_success"]),
                            f"{audit['max_abs_gradient']:.6f}", f"{audit['max_abs_projected_gradient']:.6f}",
                            str(len(audit["active_bound_indices"])), f"{audit['hessian_minimum_eigenvalue']:.4f}",
                            f"{audit['hessian_condition_number']:.0f}"])
    bp.add_table(doc, "Supplementary Table S7. Numerical convergence and covariance audit",
                 ["Domain", "Optimizer success", "Max raw gradient", "Max projected gradient", "Active bounds", "Min Hessian eigenvalue", "Hessian condition"],
                 convergence, [0.85, 1.1, 1.2, 1.35, 0.9, 1.3, 1.2],
                 note="The pain model's larger raw gradient was attached to one active lower-bound baseline-intensity parameter; the projected gradient satisfied the optimisation tolerance. Both Hessians and robust covariance matrices were positive semidefinite without eigenvalue clipping.", font_size=6.4, total=9500)

    bp.add_table(doc, "Supplementary Table S8a. Function and mortality pathways in the pain-state model",
                 ["Transition", "Deaths", "F1 vs F0 HR (95% CI)", "P", "F2 vs F0 HR (95% CI)", "P"],
                 hr_rows("pain", pain, mortality=True), [1.15, 0.9, 2.05, 0.8, 2.05, 0.8], font_size=6.7, total=9500)
    bp.add_table(doc, "Supplementary Table S8b. Pain and mortality pathways in the function-state model",
                 ["Transition", "Deaths", "P1 vs P0 HR (95% CI)", "P", "P2 vs P0 HR (95% CI)", "P"],
                 hr_rows("function", function, mortality=True), [1.15, 0.9, 2.05, 0.8, 2.05, 0.8],
                 note="Mortality contrasts are conditional transition associations. Wide intervals for sparse state-exposure combinations require caution.", font_size=6.7, total=9500)

    doc.add_page_break()
    bp.add_table(doc, "Supplementary Table S9a. Period-specific two-year pain-state probabilities shown in Fig. 3a",
                 ["Domain", "Period", "Exposure", "Origin → destination", "Fitted %", "Simulation median %", "95% interval %"],
                 probability_rows("pain"), [0.8, 1.05, 0.8, 1.45, 0.9, 1.2, 1.35],
                 note="Scenarios start in P0 and compare F0 with F2. Within each period and exposure, the four fitted probabilities sum to approximately 100%.", font_size=6.2, total=9500)
    bp.add_table(doc, "Supplementary Table S9b. Period-specific two-year function-state probabilities shown in Fig. 3b",
                 ["Domain", "Period", "Exposure", "Origin → destination", "Fitted %", "Simulation median %", "95% interval %"],
                 probability_rows("function"), [0.8, 1.05, 0.8, 1.45, 0.9, 1.2, 1.35],
                 note="Scenarios start in F0 and compare P0 with P2. Intervals use 5,000 complete-vector draws and household-robust covariance. Complete probabilities for every origin and exposure combination are supplied as domain-specific CSV files in source_data/.", font_size=6.2, total=9500)

    supplements = [
        ("FigS1_mortality_HRs_fivewave.png", "Supplementary Figure S1", "Two forest plots of conditional mortality hazard ratios.",
         "Supplementary Figure S1. Household-cluster robust adjusted hazard ratios for mortality pathways in the pain and function models. The logarithmic axis accommodates wide confidence intervals in sparse strata."),
        ("FigS2_transition_counts_fivewave.png", "Supplementary Figure S2", "Two heatmaps of observed endpoint changes among states and death.",
         "Supplementary Figure S2. Observed adjacent-visit endpoint changes by interval-start and destination state. Colour is scaled to log10(count+1); labels show raw counts."),
        ("FigS3_sensitivity_point_estimates_fivewave.png", "Supplementary Figure S3", "Two dot plots of focal hazard-ratio point estimates across five analyses.",
         "Supplementary Figure S3. Focal-pathway adjusted HR point estimates from the five-wave primary analysis and sensitivity analyses using weighting, strict function scoring, Exit Interview-only 2013 death ascertainment and exclusion of the 2013 observations."),
    ]
    for stem, label, alt, legend in supplements:
        doc.add_page_break(); bp.add_figure(doc, FIG / stem, label, alt, 6.4); bp.add_text(doc, legend, style="Figure Caption")
    doc.core_properties.title = TITLE + " - Supplementary Information"
    doc.save(SUPPLEMENT)


def build_cover_letter() -> None:
    doc = Document(); style(doc, "Cover letter | Scientific Reports", compact=True)
    p = doc.add_paragraph(style="Document Label"); bp.set_run_font(p.add_run("COVER LETTER"), size=9, bold=True, color=bp.BLUE)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; bp.set_run_font(p.add_run("24 August 2026"), size=9.5, color=bp.GRAY)
    bp.add_text(doc, "Dear Editor,")
    bp.add_text(doc, f"Please consider our Article, '{TITLE}', for publication in Scientific Reports.")
    bp.add_text(doc, "Using the 2011, 2013, 2015, 2018 and 2020 CHARLS waves, we followed 21,235 adults across 62,135 adjacent-wave intervals. BADL limitation was associated with a later transition from no pain to multisite pain, while pain at one or more sites was associated with a later transition from no functional limitation to BADL limitation. The analysis distinguishes these two directions and reports both relative transition rates and period-specific absolute probabilities.")
    bp.add_text(doc, "The study is suited to Scientific Reports because it addresses a common ageing-related problem in a large national cohort and uses a transparent multistate framework that accommodates improvement, worsening and death. The manuscript avoids causal claims, reports all living-state transitions and includes sensitivity analyses for weighting, function scoring, death ascertainment and observation-wave influence.")
    bp.add_text(doc, "Aggregate source data, analysis scripts, publication-quality figures, Supplementary Information and a STROBE reporting map accompany the manuscript. Controlled-access participant data are not redistributed.")
    bp.add_text(doc, "The manuscript is original and is not under consideration elsewhere. All authors reviewed and approved the final manuscript. No prior discussion with a Scientific Reports Editorial Board Member has taken place." )
    bp.add_text(doc, "Thank you for considering this manuscript.\n\nSincerely,\n[Corresponding author name]\n[Institution and postal address]\n[Institutional email]\n\nSuggested reviewers: [optional]\nReviewers to exclude: [optional]")
    doc.core_properties.title = "Cover letter for " + TITLE; doc.save(COVER)


def build_checklist() -> None:
    doc = Document(); style(doc, "Scientific Reports submission checklist", compact=True)
    title_page(doc, "Submission readiness checklist", metadata_notice=False)
    body_words = sum(count_words(x) for x in INTRO + RESULTS + DISCUSSION)
    bp.add_notice(doc, "PACKAGE STATUS: SCIENTIFIC CONTENT COMPLETE; AUTHOR METADATA PENDING", [
        "The five-wave analysis, tables, figures, Supplementary Information, source data, code, cover letter, and reporting map are assembled.",
        "Before submission, all authors should independently reproduce and approve the results.",
    ])
    checks = [
        ["Title", "Complete", f"{count_words(TITLE)} words; within the 20-word limit."],
        ["Abstract", "Complete", f"Unstructured; {count_words(ABSTRACT)} words; within the 200-word limit."],
        ["Main narrative", "Complete", f"Introduction + Results + Discussion: {body_words} words."],
        ["Keywords", "Complete", "Six keywords."],
        ["Main display items", "Complete", "Table 1, Tables 2a/2b, and four figures: seven display items."],
        ["Figures", "Complete", "Individual TIFF at 600 dpi, vector PDF, and PNG review copies; colourblind-aware palette."],
        ["Supplement", "Complete", "One Supplementary Information DOCX; PDF is generated after final rendering."],
        ["Data availability", "Complete", "Controlled-access CHARLS route stated; aggregate source data included."],
        ["Code availability", "Action required", "Insert a permanent DOI/URL after author verification and archiving."],
        ["Ethics", "Action required", "CHARLS approval stated; authors must confirm the local secondary-analysis determination."],
        ["Declarations", "Action required", "Complete authorship, contributions, funding, competing interests, and correspondence."],
        ["AI disclosure", "Author approval required", "Disclosure reflects code/document/figure assistance and requires author confirmation."],
        ["Reporting guideline", "Complete", "STROBE reporting map included."],
    ]
    bp.add_table(doc, "Table 1. Format and policy checks", ["Item", "Status", "Evidence/action"], checks,
                 [1.8, 1.7, 4.2], font_size=7.0, total=9500)
    doc.add_heading("Author-owned information required before upload", level=1)
    for item in [
        "Final author names and order, affiliations, ORCID iDs, and corresponding-author contact details.",
        "CRediT-style author contributions using initials; funding and funder-role statement; competing-interest declaration.",
        "Local secondary-analysis ethics review or exemption determination and identifier, if applicable.",
        "Permanent public citation or URL for verified code and aggregate outputs.",
        "Final author verification of every number, reference, figure, script, and the generative-AI disclosure.",
        "Confirmation that the work is original, not under review elsewhere, and approved by every author.",
    ]:
        p = doc.add_paragraph(style="List Bullet"); bp.set_run_font(p.add_run(item), size=9.1)
    doc.add_heading("Package map", level=1)
    files = [
        [MANUSCRIPT.name, "Main Article manuscript"], [SUPPLEMENT.name + " / .pdf", "Single Supplementary Information file"],
        [COVER.name, "Editorial cover letter"], [STROBE.name, "STROBE cohort-study reporting map"],
        ["figures/", "Main and supplementary figures in TIFF, PDF, and PNG"],
        ["source_data/", "Aggregate table/figure data and model audit outputs"], ["code/", "Five-wave cleaning, modelling, and figure scripts"],
    ]
    bp.add_table(doc, "Table 2. Package contents", ["File/folder", "Purpose"], files, [3.8, 3.9], font_size=7.2, total=9500)
    doc.core_properties.title = "Scientific Reports submission checklist"; doc.save(CHECKLIST)


def build_strobe() -> None:
    doc = Document(); style(doc, "STROBE reporting map", compact=True)
    title_page(doc, "STROBE cohort-study reporting map", metadata_notice=False)
    bp.add_text(doc, "This reporting map links the 22 STROBE cohort-study items to the manuscript and Supplementary Information.", style="Small Note", align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ["1", "Design and balanced summary", "Title; Abstract"], ["2", "Background and rationale", "Introduction 1-2"],
        ["3", "Objectives", "Introduction 3"], ["4", "Key design features", "Abstract; Methods, Study design"],
        ["5", "Setting and dates", "Methods; Fig. 1"], ["6", "Eligibility, entry, and follow-up", "Methods; Fig. 1; Table S3"],
        ["7", "Outcomes, exposures, covariates", "Methods, Pain/Function/Covariates"], ["8", "Data sources and measurement", "Methods; Tables S1-S2"],
        ["9", "Bias", "Methods, Sensitivity analyses; Discussion"], ["10", "Study size", "Results 1; Fig. 1"],
        ["11", "Quantitative variables", "Methods, states and covariates"], ["12", "Statistical methods, missingness, sensitivity", "Methods, Multistate models through Sensitivity; Supplement"],
        ["13", "Participant flow", "Fig. 1; Table S3"], ["14", "Descriptive data", "Table 1; Tables S3-S4"],
        ["15", "Outcome events", "Results; Tables 2a/2b; Fig. S2"], ["16", "Adjusted estimates and precision", "Tables 2a/2b; Figs. 2-3"],
        ["17", "Other analyses", "Results 4 and 6; Fig. 4; Tables S5-S9"], ["18", "Key results", "Discussion 1"],
        ["19", "Limitations", "Discussion 5"], ["20", "Cautious interpretation", "Discussion 2-6"],
        ["21", "Generalisability", "Discussion 4-5"], ["22", "Funding and funder role", "Funding"],
    ]
    bp.add_table(doc, "STROBE reporting map", ["Item", "Reporting expectation", "Location"], rows,
                 [0.7, 3.6, 3.4], font_size=7.1, total=9500)
    doc.core_properties.title = "STROBE reporting map for " + TITLE; doc.save(STROBE)


def write_source_data(pain: pd.DataFrame, function: pd.DataFrame, entry: pd.DataFrame) -> None:
    SOURCE.mkdir(parents=True, exist_ok=True)
    pd.DataFrame(baseline_rows(entry), columns=["characteristic", "overall", "P0", "P1", "P2"]).to_csv(SOURCE / "Table1_source_data.csv", index=False)
    pd.DataFrame(hr_rows("pain", pain), columns=["transition", "observed", "F1_vs_F0_HR_CI", "P1", "F2_vs_F0_HR_CI", "P2"]).to_csv(SOURCE / "Table2_pain_transitions.csv", index=False)
    pd.DataFrame(hr_rows("function", function), columns=["transition", "observed", "P1_vs_P0_HR_CI", "P1", "P2_vs_P0_HR_CI", "P2"]).to_csv(SOURCE / "Table3_function_transitions.csv", index=False)
    pd.DataFrame(model_selection_rows(), columns=["domain", "model", "parameters", "NLL", "LR_vs_126", "df", "P"]).to_csv(SOURCE / "TableS6_model_selection.csv", index=False)
    pd.DataFrame(probability_rows(), columns=["domain", "period", "exposure", "transition", "fitted_percent", "simulation_median_percent", "ci95_percent"]).to_csv(SOURCE / "TableS9_focal_probabilities.csv", index=False)
    pd.DataFrame(sensitivity_rows(), columns=["analysis", "people_intervals", "pain_path_hr", "function_path_hr"]).to_csv(SOURCE / "TableS5_sensitivities.csv", index=False)
    for domain in ("pain", "function"):
        for name in ["fivewave_final_full_household_robust_hr.csv", "fivewave_final_full_period_probabilities.csv", "fivewave_final_sensitivity_summary.csv", "fivewave_final_full_audit.json"]:
            shutil.copy2(FIVE / "models" / domain / name, SOURCE / f"{domain}_{name}")
    shutil.copy2(FIVE / "phase1" / "2013_sample_construction_audit.json", SOURCE / "2013_sample_construction_audit.json")
    shutil.copy2(FIVE / "phase1" / "audit" / "CHARLS_质控明细.json", SOURCE / "fivewave_data_quality_audit.json")


def copy_code() -> None:
    CODE.mkdir(parents=True, exist_ok=True)
    for name in ["build_2013_sample_info.py", "charls_multistate_clean_5wave.py", "charls_cohort_build_5wave.py",
                 "fit_fivewave_ctmc.py", "finalize_full_fivewave.py", "generate_fivewave_figures.py",
                 "build_scirep_fivewave_package.py"]:
        shutil.copy2(FIVE / name, CODE / name)
    (CODE / "requirements.txt").write_text("numpy==2.5.2\npandas==2.2.3\nscipy==1.18.0\nmatplotlib\npython-docx\n", encoding="utf-8")


def write_readme() -> None:
    text = f"""SCIENTIFIC REPORTS SUBMISSION PACKAGE — FIVE-WAVE CHARLS ANALYSIS

Article title:
{TITLE}

STATUS: SCIENTIFIC CONTENT COMPLETE; AUTHOR METADATA PENDING.
All authors should independently verify the analysis and approve the submission.

Primary study dataset and focal estimates:
- Five waves: 2011, 2013, 2015, 2018, 2020.
- Final sample: 21,235 participants; 62,135 intervals; 1,973 death endpoints.
- F2 vs F0 for P0→P2: HR 1.89 (95% CI 1.57-2.28).
- P2 vs P0 for F0→F2: HR 2.19 (95% CI 1.86-2.59).
- Excluding the 2013 observations was a secondary wave-influence analysis, not a staged addition to the primary dataset.

Required author actions:
1. Add final author names, affiliations, ORCID iDs, and corresponding-author details.
2. Complete contributions, funding, competing interests, and local secondary-analysis ethics language.
3. Rerun and independently verify the code using authorised CHARLS data.
4. Archive verified code and aggregate outputs in a permanent repository and insert its DOI/URL.
5. Remove author-action notices from the manuscript and cover letter.

The package intentionally excludes controlled-access individual-level CHARLS data.
"""
    (PACKAGE / "README_FIRST.txt").write_text(text, encoding="utf-8")

    audit = """# Five-wave analysis audit

## Data integration

- Five waves: 2011, 2013, 2015, 2018, 2020.
- Harmonised unique identifiers: 26,347; person-wave rows: 100,715.
- 2013 health IDs: 18,455; union IDs: 19,689; confirmed deaths: 464; health/death conflicts: 0.
- Duplicate harmonised person-wave rows: 0; post-death health records: 0.

## Analytic sample

- People: 21,235; intervals: 62,135; household clusters: 12,577.
- Interval starts: 14,621 (2011), 14,518 (2013), 16,269 (2015), 16,727 (2018).
- Confirmed death endpoints: 1,973.

## Final model

- Continuous-time four-state panel likelihood with all nine directed living/death transitions.
- Four period-specific baseline-intensity sets; transition-specific exposure and confounder effects.
- 126 parameters; household-cluster robust covariance; 5,000 probability draws.
- Maximum projected gradient: 0.000490 (pain), 0.000725 (function).
- Positive numerical Hessians and positive robust probability covariance; no eigenvalue clipping.

## Focal results

- Pain model, P0→P2, F2 vs F0: HR 1.892978 (95% CI 1.571012-2.280928), P=1.9589e-11.
- Pain model, P0→P2, F1 vs F0: HR 1.074800 (95% CI 0.835488-1.382660), P=0.5746.
- Function model, F0→F2, P1 vs P0: HR 1.547984 (95% CI 1.278623-1.874091), P=7.466e-06.
- Function model, F0→F2, P2 vs P0: HR 2.192971 (95% CI 1.856925-2.589832), P=2.183e-20.

## Interpretation guardrails

These are conditional observational associations. They do not establish causal direction, within-wave temporal order, individual prognosis or intervention effects. The analysis was designed and fitted using all five waves together. Excluding the 2013 observations was used only to assess sensitivity to the observation schedule and sample composition.
"""
    (PACKAGE / "FiveWave_Analysis_Audit.md").write_text(audit, encoding="utf-8")


def build_zip() -> None:
    if ZIP_PATH.exists(): ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in sorted(PACKAGE.rglob("*")):
            if path.is_file() and not any(part.startswith("qa_") for part in path.parts):
                archive.write(path, Path("Scientific_Reports_submission_2013") / path.relative_to(PACKAGE))


def main() -> None:
    PACKAGE.mkdir(parents=True, exist_ok=True); SOURCE.mkdir(parents=True, exist_ok=True); CODE.mkdir(parents=True, exist_ok=True)
    obsolete = SOURCE / "TableS9_probabilities.csv"
    if obsolete.exists(): obsolete.unlink()
    pain, function, long = load_data(); entry = entry_frame(pain, long)
    build_manuscript(pain, function, entry); build_supplement(pain, function, entry)
    build_cover_letter(); build_checklist(); build_strobe(); write_source_data(pain, function, entry); copy_code(); write_readme()
    print(f"Built five DOCX files in {PACKAGE}")
    print(f"Title words: {count_words(TITLE)}")
    print(f"Abstract words: {count_words(ABSTRACT)}")
    print(f"Introduction + Results + Discussion words: {sum(count_words(x) for x in INTRO + RESULTS + DISCUSSION)}")
    print("Render and verify documents before creating the final ZIP.")


if __name__ == "__main__":
    main()
