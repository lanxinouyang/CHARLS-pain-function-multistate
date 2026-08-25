#!/usr/bin/env python3
"""Word-formatting helpers used by the five-wave manuscript builder."""

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from table_geometry import apply_table_geometry, column_widths_from_weights

NAVY = "183B56"
BLUE = "3B82A0"
TEAL = "2A9D8F"
VERMILLION = "C94C2C"
GRAY = "697386"
PALE = "F3F7FA"
MID = "DCE8EF"
INK = "17212B"
NOTICE = "FFF4DF"


def set_run_font(run, size=None, bold=None, italic=None, color=None, font="Arial"):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rfonts.set(qn(attr), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color="D5DCE3", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def keep_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    set_run_font(run, size=8.5, color=GRAY)


def add_bottom_rule(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "D5DCE3")
    pbdr.append(bottom)
    ppr.append(pbdr)


def style_document(doc: Document, running_title: str, compact=False):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82 if compact else 0.9)
    section.bottom_margin = Inches(0.82 if compact else 0.9)
    section.left_margin = Inches(0.85 if compact else 1.0)
    section.right_margin = Inches(0.85 if compact else 1.0)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.8 if compact else 10.5)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        normal._element.rPr.rFonts.set(qn(attr), "Arial")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(5.5)

    for name, size, before, after in [("Heading 1", 14, 15, 7), ("Heading 2", 11, 9, 4)]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(INK)
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia"):
            style._element.rPr.rFonts.set(qn(attr), "Arial")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    custom = [
        ("Article Title", 18, True, INK, WD_ALIGN_PARAGRAPH.CENTER, 1.05),
        ("Author Line", 10.5, False, INK, WD_ALIGN_PARAGRAPH.CENTER, 1.1),
        ("Document Label", 8.5, True, BLUE, WD_ALIGN_PARAGRAPH.CENTER, 1.0),
        ("Table Caption", 9, True, NAVY, WD_ALIGN_PARAGRAPH.LEFT, 1.05),
        ("Figure Caption", 9, False, INK, WD_ALIGN_PARAGRAPH.LEFT, 1.1),
        ("Small Note", 8.2, False, GRAY, WD_ALIGN_PARAGRAPH.LEFT, 1.05),
        ("Reference Text", 8.3, False, INK, WD_ALIGN_PARAGRAPH.LEFT, 1.05),
        ("Checklist Status", 9, True, NAVY, WD_ALIGN_PARAGRAPH.LEFT, 1.05),
    ]
    for name, size, bold, color, align, spacing in custom:
        style = doc.styles[name] if name in doc.styles else doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia"):
            style._element.rPr.rFonts.set(qn(attr), "Arial")
        style.paragraph_format.alignment = align
        style.paragraph_format.line_spacing = spacing
        style.paragraph_format.space_after = Pt(4)
        if "Caption" in name:
            style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run(running_title), size=8.2, color=GRAY)
    add_bottom_rule(header)
    add_page_number(section.footer.paragraphs[0])


def add_text(doc, text, *, label=None, style=None, align=None, keep=False):
    p = doc.add_paragraph(style=style)
    if label and text.startswith(label):
        set_run_font(p.add_run(label), bold=True, color=INK)
        set_run_font(p.add_run(text[len(label):]))
    else:
        set_run_font(p.add_run(text))
    if align is not None:
        p.alignment = align
    if keep:
        p.paragraph_format.keep_with_next = True
    return p


def add_notice(doc, title, lines):
    table = doc.add_table(rows=1, cols=1)
    repeat_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, NOTICE)
    set_cell_borders(cell, color="D6A442", size="6")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(title), size=10, bold=True, color=VERMILLION)
    for line in lines:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(line), size=8.8, color=INK)
    apply_table_geometry(table, [9000], table_width_dxa=9000, indent_dxa=120)
    doc.add_paragraph()


def add_table(doc, caption, columns, rows, weights, note=None, font_size=7.6, total=9000):
    p = doc.add_paragraph(style="Table Caption")
    set_run_font(p.add_run(caption), size=9, bold=True, color=NAVY)
    table = doc.add_table(rows=1, cols=len(columns))
    hdr = table.rows[0]
    repeat_header(hdr)
    for j, value in enumerate(columns):
        cell = hdr.cells[j]
        cell.text = str(value)
        set_cell_shading(cell, MID)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(1.5)
            para.paragraph_format.space_before = Pt(1.5)
            para.paragraph_format.line_spacing = 1.0
            for run in para.runs:
                set_run_font(run, size=font_size, bold=True, color=NAVY)
    for i, row_values in enumerate(rows):
        row = table.add_row()
        keep_row(row)
        for j, value in enumerate(row_values):
            cell = row.cells[j]
            cell.text = str(value)
            set_cell_borders(cell)
            if i % 2:
                set_cell_shading(cell, "F8FAFC")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_after = Pt(1.2)
                para.paragraph_format.space_before = Pt(1.2)
                para.paragraph_format.line_spacing = 1.0
                for run in para.runs:
                    set_run_font(run, size=font_size)
    widths = column_widths_from_weights(weights, total)
    apply_table_geometry(table, widths, table_width_dxa=total, indent_dxa=80,
                         cell_margins_dxa={"top": 55, "bottom": 55, "start": 80, "end": 80})
    if note:
        p = doc.add_paragraph(style="Small Note")
        set_run_font(p.add_run(note), size=8.0, color=GRAY)
        p.paragraph_format.space_after = Pt(7)
    return table


def add_figure(doc, image_path, short_label, alt_text, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    picture = p.add_run().add_picture(str(image_path), width=Inches(width))
    picture._inline.docPr.set("title", short_label)
    picture._inline.docPr.set("descr", alt_text)
    cap = doc.add_paragraph(style="Figure Caption")
    set_run_font(cap.add_run(short_label), size=9, bold=True, color=NAVY)


REFERENCES = [
    "World Health Organization. Musculoskeletal health. https://www.who.int/news-room/fact-sheets/detail/musculoskeletal-conditions (2022).",
    "GBD 2021 Other Musculoskeletal Disorders Collaborators. Global, regional, and national burden of other musculoskeletal disorders, 1990-2020, and projections to 2050. Lancet Rheumatol. 5, e670-e682 (2023).",
    "Chen, H. et al. The contributions of diseases to disability burden among the elderly population in China. J. Aging Health 26, 261-282 (2014).",
    "Yu, Z. et al. Musculoskeletal disorder burden and its attributable risk factors in China: estimates and predicts from 1990 to 2044. Int. J. Environ. Res. Public Health 20, 840 (2023).",
    "Qiu, Y. et al. The prevalence and economic burden of pain on middle-aged and elderly Chinese people: results from the China Health and Retirement Longitudinal Study. BMC Health Serv. Res. 20, 600 (2020).",
    "Yu, T. et al. Assessing pain among Chinese elderly-China Health and Retirement Longitudinal Study. Iran. J. Public Health 47, 553-560 (2018).",
    "Buchman, A. S. et al. Musculoskeletal pain and incident disability in community-dwelling older adults. Arthritis Care Res. 62, 1287-1293 (2010).",
    "Shah, R. C. et al. Musculoskeletal pain is associated with incident mobility disability in community-dwelling elders. J. Gerontol. A Biol. Sci. Med. Sci. 66, 82-88 (2011).",
    "Eggermont, L. H. et al. Pain characteristics associated with the onset of disability in older adults: the MOBILIZE Boston Study. J. Am. Geriatr. Soc. 62, 1007-1016 (2014).",
    "Kaiho, Y. et al. Impact of pain on incident risk of disability in elderly Japanese: cause-specific analysis. Anesthesiology 126, 688-696 (2017).",
    "Yiengprugsawan, V. & Steptoe, A. Impacts of persistent general and site-specific pain on activities of daily living and physical performance. Geriatr. Gerontol. Int. 18, 1051-1057 (2018).",
    "Liu, P. et al. Association between multisite musculoskeletal pain and disability trajectories among community-dwelling older adults. Aging Clin. Exp. Res. 36, 115 (2024).",
    "Chu, J. et al. Pain status and disability in activities of daily living among older adults in China: evidence from CHARLS 2020. Pain Res. Manag. 2025, 4974163 (2025).",
    "Ding, H., Wang, K., Li, Y. & Zhao, X. Trends in disability in activities of daily living and instrumental activities of daily living among Chinese older adults from 2011 to 2018. Aging Clin. Exp. Res. 36, 27 (2024).",
    "Thielke, S. M. et al. Persistence and remission of musculoskeletal pain in community-dwelling older adults. J. Am. Geriatr. Soc. 60, 1393-1400 (2012).",
    "Lawton, M. P. & Brody, E. M. Assessment of older people: self-maintaining and instrumental activities of daily living. Gerontologist 9, 179-186 (1969).",
    "Andersen, P. K. & Keiding, N. Multi-state models for event history analysis. Stat. Methods Med. Res. 11, 91-115 (2002).",
    "Jackson, C. H. Multi-state models for panel data: the msm package for R. J. Stat. Softw. 38, 1-28 (2011).",
    "Sun, F., Zimmer, Z. & Zajacova, A. Pain and disability transitions among older Americans: the role of education. J. Pain 24, 1009-1019 (2023).",
    "Katz, S. et al. Studies of illness in the aged: the index of ADL. JAMA 185, 914-919 (1963).",
    "Gao, Q., Muniz Terrera, G., Mayston, R. & Prina, M. Multistate survival modelling of multimorbidity and transitions across health needs states and death. J. Epidemiol. Community Health 78, 212-219 (2024).",
    "Lynch, M., Bucknall, M., Jagger, C., Kingston, A. & Wilkie, R. Demographic, health, physical activity, and workplace factors are associated with lower healthy working life expectancy and life expectancy at age 50. Sci. Rep. 14, 5936 (2024).",
    "Zhao, Y. et al. Cohort profile: the China Health and Retirement Longitudinal Study (CHARLS). Int. J. Epidemiol. 43, 61-68 (2014).",
    "von Elm, E. et al. The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement. Lancet 370, 1453-1457 (2007).",
]
